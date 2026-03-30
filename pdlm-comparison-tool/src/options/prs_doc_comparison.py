import streamlit as st
import pandas as pd
import tempfile
import io
import re
import unicodedata
from docx import Document

def normalize_text(text):
    # Remove acentos, deixa minúsculo, tira espaços extras
    text = unicodedata.normalize('NFKD', text)
    text = ''.join([c for c in text if not unicodedata.combining(c)])
    return re.sub(r'\s+', ' ', text).strip().lower()

def normalize_spaces(series):
    return pd.Series(series).astype(str).apply(lambda x: re.sub(r'\s+', ' ', x.strip()))

def get_feature_ids_from_word(doc_path, col_name='PRS'):
    doc = Document(doc_path)
    feature_ids = []
    section_found = False
    section_key = normalize_text('13. Annex 1 – Matrix of PRS and TCs')
    # Busca aproximada pela seção
    for paragraph in doc.paragraphs:
        if section_key in normalize_text(paragraph.text):
            section_found = True
        elif section_found and paragraph.text.strip() == '':
            break
    if not section_found:
        st.warning("Section similar to '13. Annex 1 – Matrix of PRS and TCs' not found. Trying all tables.")
    # Busca por linhas com PRS na primeira coluna em todas as tabelas
    for table in doc.tables:
        for row in table.rows:
            first_cell = normalize_text(row.cells[0].text)
            if first_cell == col_name.lower() or first_cell == 'prs':
                # Pega todos os valores das colunas exceto a primeira
                ids = [cell.text.strip() for cell in row.cells[1:] if cell.text.strip()]
                feature_ids.extend([[id_] for id_ in ids if id_])
            else:
                # Mantém o comportamento anterior para linhas do tipo 'PRS: ...'
                value = row.cells[0].text.strip()
                match = re.match(r'PRS:\s*(.+)', value, re.IGNORECASE)
                if match:
                    id_value = match.group(1).strip()
                    if id_value:
                        feature_ids.append([id_value])
    return feature_ids

def get_sheet_with_fallback(xl, preferred_name):
    preferred_norm = normalize_text(preferred_name)
    for sheet in xl.sheet_names:
        if preferred_norm in normalize_text(sheet):
            return sheet
    st.warning(f"Worksheet similar to '{preferred_name}' not found. Please select the sheet that you want to analyze.")
    return st.selectbox("Select the worksheet that you want to analyze:", xl.sheet_names, key=preferred_name)

def get_urs_ids_from_excel(excel_path, sheet_name='Labeling and Learning Materials', col_name='PRS ID'):
    xl = pd.ExcelFile(excel_path)
    sheet_to_use = get_sheet_with_fallback(xl, sheet_name)
    df = pd.read_excel(excel_path, sheet_name=sheet_to_use)
    col_norm = normalize_text(col_name)
    matching_columns = [col for col in df.columns if col_norm in normalize_text(col) or 'prs' in normalize_text(col)]
    if not matching_columns:
        st.warning(f"Column similar to '{col_name}' not found in the sheet '{sheet_to_use}'. Showing available columns: {list(df.columns)}")
        return []
    return df[matching_columns[0]].dropna().astype(str).str.strip().drop_duplicates().tolist()

def pad_list(lst, target_len):
    lst = lst if lst is not None else []
    if target_len is None or not isinstance(target_len, int) or target_len < 1:
        target_len = 1
    return lst + [""] * max(target_len - len(lst), 0)

def safe_str_list(lst):
    # Garante que cada elemento é string simples, nunca método, objeto estranho, lista ou slice
    out = []
    for x in lst:
        # Só aceita string, int ou float
        if isinstance(x, (str, int, float)):
            s = str(x)
            if (
                "method" in s
                or "descriptor" in s
                or s.startswith("<")
                or s.startswith("[")
                or s.startswith("(")
                or s.strip() == ""
            ):
                continue
            out.append(s)
        # Se não for, ignora
    return out

def remove_duplicates(seq):
    seen = set()
    result = []
    for item in seq:
        if item not in seen and item != "":
            seen.add(item)
            result.append(item)
    return result

def get_psre_ids_from_excel(excel_path, sheet_name='Evaluation', col_letter='G'):
    xl = pd.ExcelFile(excel_path)
    sheet_to_use = get_sheet_with_fallback(xl, sheet_name)
    df = pd.read_excel(excel_path, sheet_name=sheet_to_use, header=None)
    col_idx = 6  # Coluna G = índice 6 (0-based)
    if df.shape[1] <= col_idx:
        st.warning(f"Aba '{sheet_to_use}' não possui coluna G (índice 6). Número de colunas encontradas: {df.shape[1]}")
        return []
    values = df.iloc[1:, col_idx].dropna().astype(str)
    ids = []
    for val in values:
        # Extrai todos os padrões TASY_PRS_ID_... da célula
        found = re.findall(r'TASY_PRS_ID_[\d\.]+', val)
        ids.extend([x.strip() for x in found if x.strip()])
    return ids

def get_rmm_ids_from_excel(excel_path, sheet_name='Risk Management Matrix', col_letter='AC'):
    xl = pd.ExcelFile(excel_path)
    sheet_to_use = get_sheet_with_fallback(xl, sheet_name)
    df = pd.read_excel(excel_path, sheet_name=sheet_to_use, header=None)
    col_idx = 28  # Coluna AC = índice 28 (0-based)
    if df.shape[1] <= col_idx:
        st.warning(f"Aba '{sheet_to_use}' não possui coluna AC (índice 28). Número de colunas encontradas: {df.shape[1]}")
        return []
    values = df.iloc[3:, col_idx].dropna().astype(str)  # pula as 3 primeiras linhas
    ids = []
    for val in values:
        found = re.findall(r'TASY_PRS_ID_[\d\.]+', val)
        ids.extend([x.strip() for x in found if x.strip()])
    return ids

def run_comparison():
    st.markdown("### PRS DOC Comparison")
    st.info("Upload your PRS DOC, Labeling Specification, PSRE DOC and RMM DOC files to compare requirements.")
    col1, col2, col3 = st.columns(3)
    with col1:
        word_file = st.file_uploader("Upload Labeling Specification (.docx)", type=["docx"], key="lb_doc")
    with col2:
        excel_file = st.file_uploader("Upload PRS DOC (.xlsx)", type=["xlsx"], key="prs_doc")
    with col3:
        psre_file = st.file_uploader("Upload PSRE DOC (.xlsx)", type=["xlsx"], key="psre_doc")
        rmm_file = st.file_uploader("Upload RMM DOC (.xlsx)", type=["xlsx"], key="rmm_doc")
    if word_file and excel_file:
        if st.button("🔍 Run Comparison", key="run_urs_tm"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_word:
                tmp_word.write(word_file.read())
                word_path = tmp_word.name
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_excel:
                tmp_excel.write(excel_file.read())
                excel_path = tmp_excel.name
            # Labeling x PRS DOC
            feature_rows = get_feature_ids_from_word(word_path)
            def is_valid_id(val):
                val = val.strip()
                return val.startswith("TASY_PRS_ID")
            feature_ids = [(row[0].strip(), row[0].strip().lower()) for row in feature_rows if row and len(row) > 0 and is_valid_id(row[0])]
            prs_ids_raw = normalize_spaces(get_urs_ids_from_excel(excel_path) or []).tolist()
            prs_ids = set(x.strip().lower() for x in prs_ids_raw)
            only_in_word = [orig for orig, norm in feature_ids if norm not in prs_ids]
            only_in_excel = [x for x in prs_ids_raw if x.strip().lower() not in set(norm for _, norm in feature_ids)]
            common = [orig for orig, norm in feature_ids if norm in prs_ids]
            only_in_word = remove_duplicates(only_in_word)
            only_in_excel = remove_duplicates(only_in_excel)
            common = remove_duplicates(common)
            max_len = max(len(only_in_word), len(only_in_excel), len(common))
            only_in_word = pad_list(only_in_word, max_len)
            only_in_excel = pad_list(only_in_excel, max_len)
            common = pad_list(common, max_len)
            data = {
                'Only in Labeling Specification DOC (PRS)': only_in_word,
                'Only in PRS DOC, tab Labeling and Learning Materials (PRS)': only_in_excel,
                'Common (Labeling ∩ PRS DOC, tab Labeling and Learning Materials)': common
            }
            summary_df = pd.DataFrame(data)
            st.dataframe(summary_df, use_container_width=True)
            buffer = io.BytesIO()
            summary_df.to_excel(buffer, index=False)
            buffer.seek(0)
            st.success("Comparison complete! Download your results below.")
            st.download_button(
                "⬇️ Download Result as Excel",
                data=buffer,
                file_name="Labeling_vs_PRS_comparison_summary.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            # PSRE x PRS DOC (Security and Privacy)
            if psre_file is not None:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_psre:
                    tmp_psre.write(psre_file.read())
                    psre_path = tmp_psre.name
                psre_ids = get_psre_ids_from_excel(psre_path)
                prs_sec_ids_raw = normalize_spaces(get_urs_ids_from_excel(excel_path, sheet_name='Security and Privacy', col_name='PRS ID') or []).tolist()
                prs_sec_ids = set(x.strip().lower() for x in prs_sec_ids_raw)
                psre_ids_norm = [(x, x.strip().lower()) for x in psre_ids]
                only_in_psre = [orig for orig, norm in psre_ids_norm if norm not in prs_sec_ids]
                only_in_prs_sec = [x for x in prs_sec_ids_raw if x.strip().lower() not in set(norm for _, norm in psre_ids_norm)]
                common_psre = [orig for orig, norm in psre_ids_norm if norm in prs_sec_ids]
                only_in_psre = remove_duplicates(only_in_psre)
                only_in_prs_sec = remove_duplicates(only_in_prs_sec)
                common_psre = remove_duplicates(common_psre)
                max_len2 = max(len(only_in_psre), len(only_in_prs_sec), len(common_psre))
                only_in_psre = pad_list(only_in_psre, max_len2)
                only_in_prs_sec = pad_list(only_in_prs_sec, max_len2)
                common_psre = pad_list(common_psre, max_len2)
                data2 = {
                    'Only in PSRE DOC (PRS)': only_in_psre,
                    'Only in PRS DOC, tab: "Security and Privacy" (PRS)': only_in_prs_sec,
                    'Common (PSRE ∩ PRS DOC tab: "Security and Privacy")': common_psre
                }
                summary_df2 = pd.DataFrame(data2)
                st.dataframe(summary_df2, use_container_width=True)
                buffer2 = io.BytesIO()
                summary_df2.to_excel(buffer2, index=False)
                buffer2.seek(0)
                st.success("PSRE Comparison complete! Download your results below.")
                st.download_button(
                    "⬇️ Download PSRE Result as Excel",
                    data=buffer2,
                    file_name="PSRE_vs_PRS_Security_comparison_summary.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            # RMM x PRS DOC (Risk Management Matrix)
            if rmm_file is not None:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_rmm:
                    tmp_rmm.write(rmm_file.read())
                    rmm_path = tmp_rmm.name
                rmm_ids = get_rmm_ids_from_excel(rmm_path)
                prs_rmm_ids_raw = normalize_spaces(get_urs_ids_from_excel(excel_path, sheet_name='Risk Management Matrix', col_name='PRS ID') or []).tolist()
                prs_rmm_ids = set(x.strip().lower() for x in prs_rmm_ids_raw)
                rmm_ids_norm = [(x, x.strip().lower()) for x in rmm_ids]
                only_in_rmm = [orig for orig, norm in rmm_ids_norm if norm not in prs_rmm_ids]
                only_in_prs_rmm = [x for x in prs_rmm_ids_raw if x.strip().lower() not in set(norm for _, norm in rmm_ids_norm)]
                common_rmm = [orig for orig, norm in rmm_ids_norm if norm in prs_rmm_ids]
                only_in_rmm = remove_duplicates(only_in_rmm)
                only_in_prs_rmm = remove_duplicates(only_in_prs_rmm)
                common_rmm = remove_duplicates(common_rmm)
                max_len3 = max(len(only_in_rmm), len(only_in_prs_rmm), len(common_rmm))
                only_in_rmm = pad_list(only_in_rmm, max_len3)
                only_in_prs_rmm = pad_list(only_in_prs_rmm, max_len3)
                common_rmm = pad_list(common_rmm, max_len3)
                data3 = {
                    'Only in RMM DOC (PRS)': only_in_rmm,
                    'Only in PRS DOC, tab: "Risk Management Matrix" (PRS)': only_in_prs_rmm,
                    'Common (RMM ∩ PRS DOC tab: "Risk Management Matrix")': common_rmm
                }
                summary_df3 = pd.DataFrame(data3)
                st.dataframe(summary_df3, use_container_width=True)
                buffer3 = io.BytesIO()
                summary_df3.to_excel(buffer3, index=False)
                buffer3.seek(0)
                st.success("RMM Comparison complete! Download your results below.")
                st.download_button(
                    "⬇️ Download RMM Result as Excel",
                    data=buffer3,
                    file_name="RMM_vs_PRS_Risk_Management_comparison_summary.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )