# BLOCO ANTI-WARNINGS
import os
os.environ['PYTHONWARNINGS'] = 'ignore'
import warnings
warnings.filterwarnings("ignore")

# Importação dos módulos principais
import streamlit as st
import pandas as pd
import tempfile
import re
import io
from openpyxl import load_workbook
import io
import sys
import logging
import openpyxl
from openpyxl import load_workbook

# Importação para processar arquivos .docx
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Tentar importar plotly para visualizações avançadas
try:
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

# Desabilitar todas as mensagens de aviso
warnings.filterwarnings("ignore")
logging.disable(logging.CRITICAL)
sys.stderr = io.StringIO()  # Redirecionar stderr

# Patch direto no openpyxl para desativar o aviso específico
def patch_openpyxl():
    """Patch específico para o aviso de 'Print area cannot be set to Defined name'"""
    try:
        # Encontrar o módulo workbook.py
        import openpyxl.reader.workbook as wb_module
        
        # Backup da função original
        original_warn = warnings.warn
        
        # Substituir com uma função que filtra esse aviso específico
        def filtered_warn(message, *args, **kwargs):
            if "Print area cannot be set to Defined name" not in str(message):
                original_warn(message, *args, **kwargs)
        
        # Aplicar o patch
        warnings.warn = filtered_warn
        
        print("Patch aplicado com sucesso no openpyxl!")
    except Exception as e:
        print(f"Erro ao aplicar patch: {e}")

# Aplicar o patch
patch_openpyxl()

# Função auxiliar para validar o formato de um TASY_PRS_ID
def is_valid_tasy_id(tasy_id):
    """
    Verifica se um ID tem o formato correto TASY_PRS_ID_X.Y.Z.W
    onde X, Y, Z, W são números, e rejeita explicitamente formatos problemáticos
    """
    # Rejeitar IDs problemáticos conhecidos explicitamente
    problematic_ids = ["TASY_PRS_ID_1", "TASY_PRS_ID_2", "TASY_PRS_ID_6"]
    if tasy_id in problematic_ids:
        return False
    
    if not isinstance(tasy_id, str) or not tasy_id.startswith("TASY_PRS_ID_"):
        return False
    
    # Verificar se o ID corresponde exatamente ao padrão com regex
    pattern = r'^TASY_PRS_ID_\d+\.\d+\.\d+\.\d+$'
    if not re.match(pattern, tasy_id):
        return False
    
    # Extrair a parte numérica
    numeric_part = tasy_id.replace("TASY_PRS_ID_", "")
    
    # Verificar se tem quatro partes separadas por ponto
    parts = numeric_part.split('.')
    if len(parts) != 4:
        return False
    
    # Verificar se cada parte é um número e tem pelo menos um dígito
    if not all(part.isdigit() and len(part) >= 1 for part in parts):
        return False
        
    # Não permitir IDs que são apenas um único dígito por parte
    if all(len(part) == 1 for part in parts):
        # Se todos os componentes têm apenas um dígito, verificar se não é um ID simples como 1.1.1.1
        if numeric_part in ["1.1.1.1", "2.2.2.2", "1.2.3.4", "5.6.7.8", "6.6.6.6"]:
            return False
    
    return True

# Função para extrair IDs no formato específico com ponto-e-vírgula
def extract_semicolon_ids(text):
    """Extrai IDs no formato 'TASY_PRS_ID_1.1.1.8378; TASY_PRS_ID_1.1.1.154; TASY_PRS_ID_1.1.1.194'"""
    valid_ids = set()
    
    # Primeiro, procurar por linhas inteiras que contêm múltiplos IDs
    lines = text.split('\n')
    for line in lines:
        if "TASY_PRS_ID" in line and ";" in line:
            # Esta linha pode conter IDs separados por ponto-e-vírgula
            ids_in_line = re.findall(r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+', line)
            for id in ids_in_line:
                if is_valid_tasy_id(id):
                    valid_ids.add(id)
    
    # Segundo, usar uma abordagem mais específica para capturar grupos de IDs
    pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+\s*(?:;\s*TASY_PRS_ID_\d+\.\d+\.\d+\.\d+)+'
    for match in re.finditer(pattern, text):
        chunk = match.group(0)
        ids_in_chunk = re.findall(r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+', chunk)
        for id in ids_in_chunk:
            if is_valid_tasy_id(id):
                valid_ids.add(id)
    
    return valid_ids

# Função para extrair TASY_PRS_IDs de um arquivo .docx (Risk Matrix)
def extract_tasy_ids_from_docx(file_path):
    """Extrai todos os TASY_PRS_ID de um arquivo Word .docx na seção '13. Annex 1 – Matrix of PRS and TCs'"""
    if not DOCX_AVAILABLE:
        st.error("❌ Biblioteca python-docx não está instalada. Instale com 'pip install python-docx'")
        return []
    
    all_ids = set()
    found_section = False
    
    try:
        # Abrir o documento Word
        doc = Document(file_path)
        st.write("Analisando documento Word em busca da seção '13. Annex 1 – Matrix of PRS and TCs'...")
        
        # Procurar pela seção específica
        for para in doc.paragraphs:
            if "13. Annex 1" in para.text and "Matrix of PRS" in para.text:
                found_section = True
                st.success(f"✅ Seção '13. Annex 1 – Matrix of PRS and TCs' encontrada!")
                break
        
        if not found_section:
            # Tentar buscar por apenas "Annex 1" como alternativa
            for para in doc.paragraphs:
                if "Annex 1" in para.text and "Matrix" in para.text:
                    found_section = True
                    st.success(f"✅ Seção contendo 'Annex 1' e 'Matrix' encontrada!")
                    break
        
        # Se encontramos a seção, processar todas as tabelas do documento
        if found_section:
            # Analisar todas as tabelas (ou podemos limitar às tabelas após a seção encontrada)
            prs_tables_found = 0
            
            # Para cada tabela no documento
            for table in doc.tables:
                table_has_prs = False
                prs_ids_in_table = []
                
                # Verificar se a primeira coluna de alguma linha tem "PRS"
                for row in table.rows:
                    if row.cells and len(row.cells) > 0:
                        if "PRS" in row.cells[0].text.strip():
                            table_has_prs = True
                            # Esta linha tem PRS, então extraímos os IDs de cada célula
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                # Buscar padrões de TASY_PRS_ID
                                pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
                                matches = re.findall(pattern, cell_text)
                                
                                # Filtrar para IDs válidos
                                valid_ids = [id for id in matches if is_valid_tasy_id(id)]
                                prs_ids_in_table.extend(valid_ids)
                
                if table_has_prs:
                    prs_tables_found += 1
                    all_ids.update(prs_ids_in_table)
                    st.success(f"✅ Tabela {prs_tables_found}: Encontrados {len(prs_ids_in_table)} TASY_PRS_IDs válidos")
            
            if prs_tables_found == 0:
                st.warning("⚠️ Nenhuma tabela com linha 'PRS' foi encontrada no documento.")
            else:
                st.success(f"✅ Total: {prs_tables_found} tabelas com 'PRS' processadas, {len(all_ids)} IDs únicos encontrados")
        else:
            st.warning("⚠️ Seção '13. Annex 1 – Matrix of PRS and TCs' não foi encontrada no documento.")
            
    except Exception as e:
        st.error(f"❌ Erro ao processar o arquivo Word: {str(e)}")
    
    # Se não encontrou nada, mostrar um aviso e permitir entrada manual
    if not all_ids:
        st.warning("⚠️ Nenhum TASY_PRS_ID encontrado no arquivo .docx")
        
        # Opção para entrada manual
        if st.checkbox("Deseja inserir valores manualmente?"):
            manual_input = st.text_area(
                "Cole os TASY_PRS_ID (um por linha):",
                height=200,
                placeholder="TASY_PRS_ID_6.10.601.25919\nTASY_PRS_ID_6.10.602.25904\n..."
            )
            
            if manual_input.strip():
                # Processar entrada manual
                for line in manual_input.split('\n'):
                    line = line.strip()
                    if line and 'TASY_PRS_ID_' in line and is_valid_tasy_id(line):
                        all_ids.add(line)
    
    # Exibir estatísticas finais
    if all_ids:
        st.success(f"Total final: {len(all_ids)} TASY_PRS_ID extraídos do arquivo Word")
    
    return sorted(list(all_ids))

# Função para extrair TASY_PRS_IDs de um arquivo Labeling Specification
def extract_tasy_ids_from_labeling_spec(file_path):
    """Extrai todos os TASY_PRS_ID de um arquivo Labeling Specification focando nas tabelas da seção '13. Annex 1'"""
    if not DOCX_AVAILABLE:
        st.error("❌ Biblioteca python-docx não está instalada. Instale com 'pip install python-docx'")
        return []
    
    all_ids = set()
    found_section = False
    section_index = -1
    
    try:
        # Abrir o documento Word
        doc = Document(file_path)
        st.write("Analisando Labeling Specification em busca da seção '13. Annex 1 – Matrix of PRS and TCs'...")
        
        # Procurar pela seção específica
        for i, para in enumerate(doc.paragraphs):
            if "13." in para.text and "Annex 1" in para.text and "Matrix of PRS" in para.text:
                found_section = True
                section_index = i
                st.success(f"✅ Seção '13. Annex 1 – Matrix of PRS and TCs' encontrada no parágrafo {i}!")
                break
        
        if not found_section:
            # Tentar abordagens alternativas mais flexíveis
            for i, para in enumerate(doc.paragraphs):
                if "Annex 1" in para.text and "Matrix of PRS" in para.text:
                    found_section = True
                    section_index = i
                    st.success(f"✅ Seção contendo 'Annex 1' e 'Matrix of PRS' encontrada no parágrafo {i}!")
                    break
                elif "13." in para.text and "Annex" in para.text:
                    found_section = True
                    section_index = i
                    st.success(f"✅ Seção '13. Annex' encontrada no parágrafo {i}!")
                    break
        
        # Se encontramos a seção, processar todas as tabelas do documento
        if found_section:
            # Analisar todas as tabelas que aparecem após a seção encontrada
            prs_tables_found = 0
            tables_analyzed = 0
            
            # Para cada tabela no documento
            for table in doc.tables:
                tables_analyzed += 1
                table_has_prs = False
                prs_ids_in_table = []
                
                # Verificar cada linha da tabela
                for row_idx, row in enumerate(table.rows):
                    # Verificar se a primeira célula da linha contém "PRS"
                    if row.cells and len(row.cells) > 0:
                        first_cell_text = row.cells[0].text.strip()
                        if "PRS" in first_cell_text:
                            table_has_prs = True
                            st.info(f"✓ Linha com 'PRS' encontrada na tabela {tables_analyzed}, linha {row_idx+1}")
                            
                            # Analisar todas as células da linha que contém "PRS"
                            for cell_idx, cell in enumerate(row.cells):
                                cell_text = cell.text.strip()
                                # Buscar padrões de TASY_PRS_ID
                                pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
                                matches = re.findall(pattern, cell_text)
                                
                                # Filtrar para IDs válidos
                                valid_ids = [id for id in matches if is_valid_tasy_id(id)]
                                
                                if valid_ids:
                                    st.info(f"  - Coluna {cell_idx+1}: {len(valid_ids)} IDs encontrados")
                                    prs_ids_in_table.extend(valid_ids)
                
                if table_has_prs:
                    prs_tables_found += 1
                    all_ids.update(prs_ids_in_table)
                    st.success(f"✅ Tabela {tables_analyzed}: Encontrados {len(prs_ids_in_table)} TASY_PRS_IDs válidos")
                    
                    # Mostrar os IDs encontrados nesta tabela
                    if prs_ids_in_table:
                        with st.expander(f"🔍 IDs encontrados na tabela {tables_analyzed}"):
                            for id in prs_ids_in_table:
                                st.write(f"- {id}")
            
            if prs_tables_found == 0:
                st.warning("⚠️ Nenhuma tabela com linha 'PRS' foi encontrada no Labeling Specification.")
                
                # Como fallback, buscar por IDs em todo o documento
                st.info("🔍 Tentando buscar TASY_PRS_IDs em todo o documento como alternativa...")
                
                # Buscar IDs em todos os parágrafos
                para_ids = []
                for para in doc.paragraphs:
                    para_text = para.text.strip()
                    pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
                    matches = re.findall(pattern, para_text)
                    valid_ids = [id for id in matches if is_valid_tasy_id(id)]
                    para_ids.extend(valid_ids)
                
                if para_ids:
                    all_ids.update(para_ids)
                    st.success(f"✅ Encontrados {len(para_ids)} TASY_PRS_IDs nos parágrafos do documento")
            else:
                st.success(f"✅ Total: {prs_tables_found} tabelas com 'PRS' processadas, {len(all_ids)} IDs únicos encontrados")
        else:
            st.warning("⚠️ Seção '13. Annex 1 – Matrix of PRS and TCs' não foi encontrada no Labeling Specification.")
            
            # Como fallback, buscar por qualquer tabela com "PRS" na primeira coluna
            st.info("🔍 Tentando buscar tabelas com 'PRS' em todo o documento...")
            
            prs_tables_found = 0
            tables_analyzed = 0
            
            # Para cada tabela no documento
            for table in doc.tables:
                tables_analyzed += 1
                table_has_prs = False
                prs_ids_in_table = []
                
                # Verificar cada linha da tabela
                for row in table.rows:
                    # Verificar se a primeira célula da linha contém "PRS"
                    if row.cells and len(row.cells) > 0:
                        if "PRS" in row.cells[0].text.strip():
                            table_has_prs = True
                            
                            # Analisar todas as células da linha
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                # Buscar padrões de TASY_PRS_ID
                                pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
                                matches = re.findall(pattern, cell_text)
                                
                                # Filtrar para IDs válidos
                                valid_ids = [id for id in matches if is_valid_tasy_id(id)]
                                prs_ids_in_table.extend(valid_ids)
                
                if table_has_prs:
                    prs_tables_found += 1
                    all_ids.update(prs_ids_in_table)
                    st.success(f"✅ Tabela {tables_analyzed}: Encontrados {len(prs_ids_in_table)} TASY_PRS_IDs válidos")
            
            if prs_tables_found == 0:
                st.warning("⚠️ Nenhuma tabela com linha 'PRS' foi encontrada em todo o documento.")
            else:
                st.success(f"✅ Total: {prs_tables_found} tabelas com 'PRS' processadas, {len(all_ids)} IDs únicos encontrados")
            
    except Exception as e:
        st.error(f"❌ Erro ao processar o Labeling Specification: {str(e)}")
        st.error(f"Detalhes do erro: {type(e).__name__}")
        import traceback
        st.error(traceback.format_exc())
    
    # Se não encontrou nada, mostrar um aviso e permitir entrada manual
    if not all_ids:
        st.warning("⚠️ Nenhum TASY_PRS_ID encontrado no Labeling Specification")
        
        # Opção para entrada manual
        if st.checkbox("Deseja inserir valores manualmente?", key="manual_input_labeling_spec"):
            manual_input = st.text_area(
                "Cole os TASY_PRS_ID (um por linha):",
                height=200,
                placeholder="TASY_PRS_ID_6.10.601.25919\nTASY_PRS_ID_6.10.602.25904\n...",
                key="text_area_labeling_spec"
            )
            
            if manual_input.strip():
                # Processar entrada manual
                for line in manual_input.split('\n'):
                    line = line.strip()
                    if line and 'TASY_PRS_ID_' in line and is_valid_tasy_id(line):
                        all_ids.add(line)
    
    # Exibir estatísticas finais
    if all_ids:
        st.success(f"Total final: {len(all_ids)} TASY_PRS_ID extraídos do Labeling Specification")
    
    return sorted(list(all_ids))

# Função para extrair TASY_PRS_ID de um arquivo Excel
def extract_tasy_ids(file_path):
    """Extrai todos os TASY_PRS_ID de um arquivo Excel usando múltiplos métodos"""
    all_ids = set()
    
    # Método 1: Usar pandas para ler todas as colunas/linhas da planilha específica
    try:
        st.write("Analisando planilha 'Risk Management Matrix'...")
        
        # Tentar carregar a aba específica da planilha
        try:
            # Tentar ler a aba específica "Risk Management Matrix"
            df = pd.read_excel(file_path, sheet_name="Risk Management Matrix")
            st.success("✅ Planilha 'Risk Management Matrix' encontrada!")
        except Exception as e:
            st.warning(f"Não foi possível encontrar a aba 'Risk Management Matrix': {str(e)}")
            # Tentar ler a primeira aba como fallback
            df = pd.read_excel(file_path)
            st.info("ℹ️ Usando a primeira aba da planilha como alternativa.")
        
        # Converter toda a planilha para string e concatenar em um único texto
        all_text = ""
        for column in df.columns:
            # Converter a coluna para string
            df[column] = df[column].astype(str)
            # Adicionar ao texto total
            all_text += df[column].str.cat(sep=" ") + " "
            
        # Procurar por IDs no texto combinado
        pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
        matches = re.findall(pattern, all_text)
        
        # Filtrar para IDs válidos
        valid_ids = [id for id in matches if is_valid_tasy_id(id)]
        
        # Extrair IDs separados por ponto-e-vírgula
        semicolon_ids = extract_semicolon_ids(all_text)
        
        # Adicionar todos os IDs encontrados
        all_ids.update(valid_ids)
        all_ids.update(semicolon_ids)
        
        # Mostrar estatísticas
        st.success(f"✅ Método Pandas: Encontrados {len(valid_ids)} IDs individuais e {len(semicolon_ids)} IDs em formato ponto-e-vírgula")
    
    except Exception as e:
        st.warning(f"Erro no método Pandas: {str(e)}")
    
    # Método 2: Método binário como backup - ler o arquivo como binário e procurar padrões
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            text = content.decode('utf-8', errors='ignore')
            
            # Depurar conteúdo do arquivo (para encontrar os IDs problemáticos)
            st.write("Analisando conteúdo binário do arquivo como backup...")
            
            # Procurar especificamente por padrões de IDs válidos (mais restritivo)
            # Formato TASY_PRS_ID seguido de 4 grupos de dígitos separados por pontos
            pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
            matches = re.findall(pattern, text)
            
            # Filtrar para remover IDs problemáticos conhecidos
            problematic_ids = ["TASY_PRS_ID_1", "TASY_PRS_ID_2", "TASY_PRS_ID_6"]
            valid_ids = [id for id in matches if is_valid_tasy_id(id) and id not in problematic_ids]
            
            # Mostrar todos os IDs encontrados para debug
            with st.expander("🔍 Debug - Todos os IDs encontrados"):
                st.write("IDs encontrados no arquivo:")
                for id in matches:
                    st.write(f"- {id} {'✅ Válido' if is_valid_tasy_id(id) else '❌ Inválido'}")
            
            # Adicionar os IDs válidos ao conjunto
            all_ids.update(valid_ids)
                
            # Log para depuração
            if valid_ids:
                st.success(f"✅ Encontrados {len(valid_ids)} IDs no formato padrão")
                
            # Procurar especificamente pelo padrão "TASY_PRS_ID_X.Y.Z.W; TASY_PRS_ID_A.B.C.D; ..."
            # Este padrão busca explicitamente por IDs separados por ponto e vírgula
            multiline_pattern = r'(TASY_PRS_ID_\d+\.\d+\.\d+\.\d+(?:\s*;\s*TASY_PRS_ID_\d+\.\d+\.\d+\.\d+)*)'
            multiline_matches = re.findall(multiline_pattern, text)
            
            # Depurar os blocos encontrados
            if multiline_matches:
                with st.expander("🔍 Debug - Blocos com múltiplos IDs"):
                    for i, block in enumerate(multiline_matches):
                        st.write(f"Bloco {i+1}: {block}")
            
            # Contabilizar todos os IDs válidos extraídos
            multiline_ids = []
            
            for match_group in multiline_matches:
                # Dividir por ponto e vírgula e extrair cada ID
                individual_ids = re.findall(r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+', match_group)
                
                # Mostrar os IDs encontrados neste bloco
                for tasy_id in individual_ids:
                    # Verificar se o ID tem o formato correto usando a função auxiliar
                    if is_valid_tasy_id(tasy_id):
                        all_ids.add(tasy_id)
                        multiline_ids.append(tasy_id)
                    
            if multiline_ids:
                st.success(f"✅ Encontrados {len(multiline_ids)} IDs válidos em blocos com múltiplos IDs")
                
            # Extrair especificamente de texto que segue um padrão numerado como no exemplo
            # "1. Display an error message... TASY_PRS_ID_6.10.601.25919"
            numbered_pattern = r'\d+\.\s+.*?TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
            numbered_matches = re.findall(numbered_pattern, text, re.DOTALL)
            
            # Depurar os itens numerados encontrados
            if numbered_matches:
                with st.expander("🔍 Debug - Itens numerados com IDs"):
                    for i, item in enumerate(numbered_matches[:10]):  # Mostrar apenas os primeiros 10 para não sobrecarregar
                        st.write(f"Item {i+1}: {item[:100]}...") # Mostrar apenas os primeiros 100 caracteres
            
            # Contabilizar todos os IDs válidos extraídos de itens numerados
            numbered_ids = []
            
            for numbered_text in numbered_matches:
                # Extrair todos os TASY_PRS_ID deste texto numerado
                for tasy_id in re.findall(r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+', numbered_text):
                    # Verificar o formato correto do ID usando a função auxiliar
                    if is_valid_tasy_id(tasy_id):
                        all_ids.add(tasy_id)
                        numbered_ids.append(tasy_id)
                    
            if numbered_ids:
                st.success(f"✅ Encontrados {len(numbered_ids)} IDs válidos em itens numerados")
            
            # MÉTODO ESPECÍFICO PARA O FORMATO "TASY_PRS_ID_1.1.1.8378; TASY_PRS_ID_1.1.1.154; TASY_PRS_ID_1.1.1.194"
            st.write("🔍 Buscando especificamente por IDs no formato com ponto-e-vírgula...")
            
            # Usar nossa função especializada para extrair IDs com ponto-e-vírgula
            semicolon_ids = extract_semicolon_ids(text)
            
            # Adicionar os IDs encontrados ao conjunto
            all_ids.update(semicolon_ids)
            
            # Depurar os IDs encontrados
            if semicolon_ids:
                with st.expander("🔍 Debug - IDs com ponto-e-vírgula"):
                    for i, id in enumerate(semicolon_ids):
                        st.write(f"ID {i+1}: {id}")
                
                st.success(f"✅ Encontrados {len(semicolon_ids)} IDs com ponto-e-vírgula")
    except Exception as e:
        st.warning(f"Erro na extração binária: {e}")
    
    # Se não encontrou nada, mostrar um aviso e permitir entrada manual
    if not all_ids:
        st.warning("⚠️ Nenhum TASY_PRS_ID encontrado no arquivo. Por favor, verifique o formato ou considere a entrada manual.")
        st.info("Formato esperado: TASY_PRS_ID_6.10.601.25919 ou múltiplos IDs separados por ponto e vírgula")
        
        # Opção para entrada manual, mas SEM valores padrão
        if st.checkbox("Deseja inserir valores manualmente?"):
            manual_input = st.text_area(
                "Cole os TASY_PRS_ID (um por linha):",
                height=200,
                placeholder="TASY_PRS_ID_6.10.601.25919\nTASY_PRS_ID_6.10.602.25904\n..."
            )
            
            if manual_input.strip():
                # Processar entrada manual
                for line in manual_input.split('\n'):
                    line = line.strip()
                    if line and 'TASY_PRS_ID_' in line and is_valid_tasy_id(line):
                        all_ids.add(line)
    
    # Exibir estatísticas finais
    if all_ids:
        st.success(f"Total final: {len(all_ids)} TASY_PRS_ID extraídos")
    
    return sorted(list(all_ids))

# Função para extrair TASY_PRS_IDs do PRS DOC
def extract_tasy_ids_from_prs_doc(file_path, sheet_name=None):
    try:
        # Armazenar todos os IDs encontrados
        all_ids = set()
        
        # Verificar se o arquivo tem múltiplas abas
        xl = pd.ExcelFile(file_path)
        sheet_names = xl.sheet_names
        
        st.write(f"Analisando arquivo PRS DOC com {len(sheet_names)} abas...")
        
        # Definir a aba alvo com base no parâmetro ou usar a padrão
        target_sheet = sheet_name if sheet_name else "Labeling and Learning Materials"
        if target_sheet in sheet_names:
            st.success(f"✅ Aba '{target_sheet}' encontrada!")
            try:
                # Carregar a aba específica
                df = pd.read_excel(file_path, sheet_name=target_sheet)
                
                # Verificar se existe a coluna "PRS ID"
                prs_id_column = None
                for col in df.columns:
                    if "PRS ID" in str(col):
                        prs_id_column = col
                        st.success(f"✅ Coluna '{col}' encontrada!")
                        break
                
                if prs_id_column:
                    # Extrair IDs da coluna específica
                    df[prs_id_column] = df[prs_id_column].astype(str)
                    
                    # Concatenar todos os valores da coluna
                    all_text = df[prs_id_column].str.cat(sep=" ")
                    
                    # Procurar por IDs no texto combinado
                    pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
                    matches = re.findall(pattern, all_text)
                    
                    # Filtrar para IDs válidos
                    valid_ids = [id for id in matches if is_valid_tasy_id(id)]
                    
                    # Extrair IDs separados por ponto-e-vírgula
                    semicolon_ids = extract_semicolon_ids(all_text)
                    
                    # Adicionar à coleção principal
                    all_ids.update(valid_ids)
                    all_ids.update(semicolon_ids)
                    
                    st.success(f"✅ Coluna '{prs_id_column}': Encontrados {len(valid_ids)} IDs individuais e {len(semicolon_ids)} IDs em formato ponto-e-vírgula")
                else:
                    st.warning(f"⚠️ Coluna 'PRS ID' não encontrada na aba '{target_sheet}'")
                    
                    # Mostrar colunas disponíveis para debug
                    with st.expander("🔍 Debug - Colunas disponíveis"):
                        st.write([str(col) for col in df.columns])
                    
                    # Como fallback, processar todas as colunas
                    st.info("Tentando buscar em todas as colunas como alternativa...")
                    
                    # Converter todas as colunas para um único texto
                    all_text = ""
                    for column in df.columns:
                        # Converter para string
                        df[column] = df[column].astype(str)
                        # Adicionar ao texto total
                        all_text += df[column].str.cat(sep=" ") + " "
                    
                    # Procurar por IDs no texto combinado
                    pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
                    matches = re.findall(pattern, all_text)
                    
                    # Filtrar para IDs válidos
                    valid_ids = [id for id in matches if is_valid_tasy_id(id)]
                    
                    # Extrair IDs separados por ponto-e-vírgula
                    semicolon_ids = extract_semicolon_ids(all_text)
                    
                    # Adicionar à coleção principal
                    all_ids.update(valid_ids)
                    all_ids.update(semicolon_ids)
                    
                    st.success(f"✅ Aba '{target_sheet}' (todas as colunas): Encontrados {len(valid_ids)} IDs individuais e {len(semicolon_ids)} IDs em formato ponto-e-vírgula")
            
            except Exception as e:
                st.warning(f"Erro ao processar aba '{target_sheet}': {str(e)}")
                
        else:
            st.warning(f"⚠️ Aba '{target_sheet}' não encontrada. Processando todas as abas como alternativa.")
            
            # Processar cada aba do arquivo (método original)
            for sheet_name in sheet_names:
                try:
                    # Carregar a aba atual
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Converter todas as colunas para um único texto
                    all_text = ""
                    for column in df.columns:
                        # Converter para string
                        df[column] = df[column].astype(str)
                        # Adicionar ao texto total
                        all_text += df[column].str.cat(sep=" ") + " "
                    
                    # Procurar por IDs no texto combinado
                    pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
                    matches = re.findall(pattern, all_text)
                    
                    # Filtrar para IDs válidos
                    valid_ids = [id for id in matches if is_valid_tasy_id(id)]
                    
                    # Extrair IDs separados por ponto-e-vírgula
                    semicolon_ids = extract_semicolon_ids(all_text)
                    
                    # Adicionar à coleção principal
                    all_ids.update(valid_ids)
                    all_ids.update(semicolon_ids)
                    
                    st.success(f"✅ Aba '{sheet_name}': Encontrados {len(valid_ids)} IDs individuais e {len(semicolon_ids)} IDs em formato ponto-e-vírgula")
                
                except Exception as e:
                    st.warning(f"Erro ao processar aba '{sheet_name}': {str(e)}")
        
        # Converter para lista, remover duplicatas e ordenar
        unique_ids = sorted(list(all_ids))
        
        st.success(f"✅ Total de IDs únicos encontrados no PRS DOC: {len(unique_ids)}")
        return unique_ids
    
    except Exception as e:
        st.error(f"Erro ao processar o PRS DOC: {str(e)}")
        return []

# Função para comparar TASY_PRS_IDs entre dois conjuntos
def compare_tasy_ids(ids1, ids2):
    # Converter para conjuntos para operações de conjunto
    set1 = set(ids1)
    set2 = set(ids2)
    
    # Encontrar IDs comuns
    common_ids = list(set1.intersection(set2))
    common_ids.sort()
    
    # Encontrar IDs exclusivos no primeiro conjunto (Risk Matrix)
    only_in_set1 = list(set1.difference(set2))
    only_in_set1.sort()
    
    # Encontrar IDs exclusivos no segundo conjunto (PRS DOC)
    only_in_set2 = list(set2.difference(set1))
    only_in_set2.sort()
    
    return {
        'common': common_ids,
        'only_in_rm': only_in_set1,
        'only_in_prs': only_in_set2
    }

# Função principal para executar a comparação
def run_comparison():
    st.title("PRS DOC x Documentação TASY_PRS_ID Comparison")
    st.info("Esta ferramenta extrai TASY_PRS_ID de diferentes documentos e compara com o PRS DOC")
    
    # Opção para escolher o tipo de análise
    analysis_type = st.radio(
        "Escolha o tipo de análise:",
        ["Upload de Risk Matrix", "Upload de Labeling Specification", "Entrada Direta de IDs"],
        horizontal=True
    )
    
    # Variável para controlar se é entrada direta
    use_direct_input = (analysis_type == "Entrada Direta de IDs")
    
    if use_direct_input:
        st.info("Cole o texto contendo os TASY_PRS_ID (no formato do exemplo abaixo)")
        st.markdown("""
        **Exemplo de formato aceito:**
        ```
        1. Display an error message when the database is not available
        TASY_PRS_ID_6.10.601.25919
        2. The system deployed on cloud returns a timeout message
        TASY_PRS_ID_6.10.602.25904
        3. The database shall rollback the pending transactions
        TASY_PRS_ID_6.10.601.25920
        4. Design a verifiable algorithm that calculate dosage
        TASY_PRS_ID_1.1.1.8378; TASY_PRS_ID_1.1.1.154; TASY_PRS_ID_1.1.1.194
        ```
        """)
        
        direct_input = st.text_area(
            "Cole o texto aqui:", 
            height=300,
            placeholder="1. Display an error message...\nTASY_PRS_ID_6.10.601.25919\n..."
        )
        
        if direct_input:
            with st.spinner("Processando texto..."):
                # Usar nossa função especializada para extrair IDs, incluindo os separados por ponto-e-vírgula
                semicolon_ids = extract_semicolon_ids(direct_input)
                
                # Também extrair IDs isolados
                pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
                matches = re.findall(pattern, direct_input)
                
                # Filtrar para garantir que apenas IDs com formato correto sejam incluídos
                valid_matches = [id for id in matches if is_valid_tasy_id(id)]
                
                # Combinar os dois conjuntos de IDs
                rm_ids = list(set(valid_matches).union(semicolon_ids))
                
                if rm_ids:
                    st.success(f"✅ Encontrados {len(rm_ids)} TASY_PRS_ID no texto")
                    
                    # Mostrar informações adicionais sobre o processo de extração
                    with st.expander("🔍 Debug - Processo de extração"):
                        st.write(f"- IDs válidos encontrados: {len(rm_ids)}")
                        st.write(f"- IDs com separador de ponto-e-vírgula: {len(semicolon_ids)}")
                        st.write(f"- Total de matches brutos: {len(matches)}")
                        st.write(f"- IDs inválidos descartados: {len(matches) - len(valid_matches)}")
                    
                    # Mostrar IDs em formato de tabela para melhor visualização
                    df = pd.DataFrame({"TASY_PRS_ID": sorted(rm_ids)})
                    st.dataframe(df, use_container_width=True)
                    
                    # Opção para baixar como CSV
                    csv = df.to_csv(index=False).encode('utf-8')
                    st.download_button(
                        "⬇️ Download CSV",
                        csv,
                        "tasy_prs_ids.csv",
                        "text/csv",
                        key='download-csv-direct'
                    )
                    
                    # Opção para comparar com PRS DOC
                    prs_file = st.file_uploader("Carregar PRS DOC para comparação", type=["xlsx"])
                    
                    if prs_file and st.button("Comparar com PRS DOC"):
                        with st.spinner("Extraindo TASY_PRS_ID do PRS DOC..."):
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                                tmp.write(prs_file.read())
                                prs_path = tmp.name
                            
                            # Para Risk Matrix, sempre extrair da aba padrão neste bloco
                            prs_ids = extract_tasy_ids_from_prs_doc(prs_path)
                            
                            if prs_ids:
                                st.success(f"✅ Encontrados {len(prs_ids)} TASY_PRS_ID no PRS DOC")
                                
                                # Realizar a comparação
                                comparison = compare_tasy_ids(rm_ids, prs_ids)
                                
                                # Exibir resultados da comparação
                                st.subheader("Resultados da Comparação")
                                
                                col1, col2, col3 = st.columns(3)
                                
                                with col1:
                                    st.metric("IDs em comum", len(comparison['common']))
                                
                                with col2:
                                    st.metric("Somente na Risk Matrix", len(comparison['only_in_rm']))
                                
                                with col3:
                                    st.metric("Somente no PRS DOC", len(comparison['only_in_prs']))
                                
                                # Mostrar detalhes em abas
                                tab1, tab2, tab3 = st.tabs([
                                    "🔄 IDs em comum", 
                                    "⚠️ Somente na Risk Matrix", 
                                    "❓ Somente no PRS DOC"
                                ])
                                
                                with tab1:
                                    if comparison['common']:
                                        st.dataframe(pd.DataFrame({"TASY_PRS_ID": comparison['common']}), use_container_width=True)
                                    else:
                                        st.warning("Não há IDs em comum!")
                                
                                with tab2:
                                    if comparison['only_in_rm']:
                                        st.dataframe(pd.DataFrame({"TASY_PRS_ID": comparison['only_in_rm']}), use_container_width=True)
                                    else:
                                        st.success("Todos os IDs da Risk Matrix estão no PRS DOC!")
                                
                                with tab3:
                                    if comparison['only_in_prs']:
                                        st.dataframe(pd.DataFrame({"TASY_PRS_ID": comparison['only_in_prs']}), use_container_width=True)
                                    else:
                                        st.success("Todos os IDs do PRS DOC estão na Risk Matrix!")
                                
                                # Opção para exportar resultados
                                st.subheader("Exportar Resultados")
                                
                                results = {
                                    "common_ids": comparison['common'],
                                    "only_in_rm": comparison['only_in_rm'],
                                    "only_in_prs": comparison['only_in_prs']
                                }
                                
                                # Criar um DataFrame para cada categoria
                                df_common = pd.DataFrame({"TASY_PRS_ID": comparison['common'], "Status": "Presente em ambos"})
                                df_only_rm = pd.DataFrame({"TASY_PRS_ID": comparison['only_in_rm'], "Status": "Somente na Risk Matrix"})
                                df_only_prs = pd.DataFrame({"TASY_PRS_ID": comparison['only_in_prs'], "Status": "Somente no PRS DOC"})
                                
                                # Juntar todos em um único DataFrame
                                df_all = pd.concat([df_common, df_only_rm, df_only_prs])
                                
                                # Opção para baixar como CSV
                                csv_all = df_all.to_csv(index=False).encode('utf-8')
                                st.download_button(
                                    "⬇️ Download CSV com Todos os Resultados",
                                    csv_all,
                                    "comparison_results.csv",
                                    "text/csv",
                                    key='download-csv-all'
                                )
                            else:
                                st.error("❌ Nenhum TASY_PRS_ID encontrado no PRS DOC.")
                else:
                    st.error("❌ Nenhum TASY_PRS_ID encontrado no texto. Verifique o formato.")
    elif analysis_type == "Upload de Risk Matrix":
        st.header("Análise de Risk Matrix")
        
        col1, col2 = st.columns(2)
        with col1:
            rm_file_type = st.radio(
                "Formato do arquivo Risk Matrix:",
                ["Excel (.xlsx)", "Word (.docx)"],
                horizontal=True
            )
            
            if rm_file_type == "Excel (.xlsx)":
                rm_file = st.file_uploader("Carregar Risk Management Matrix (Excel)", type=["xlsx"])
            else:
                rm_file = st.file_uploader("Carregar Risk Management Matrix (Word)", type=["docx"])
                if not DOCX_AVAILABLE and rm_file:
                    st.error("❌ Biblioteca python-docx não está instalada. Execute 'pip install python-docx' e reinicie a aplicação.")
                    rm_file = None
        
        with col2:
            prs_file = st.file_uploader("Carregar PRS DOC", type=["xlsx"])
        
        # Definir tipo de documento para uso posterior
        doc_type = "Risk Management Matrix"
            
    elif analysis_type == "Upload de Labeling Specification":
        st.header("Análise de Labeling Specification")
        
        col1, col2 = st.columns(2)
        with col1:
            st.info("📄 O arquivo Labeling Specification deve conter seção '13. Annex 1 – Matrix of PRS and TCs'")
            rm_file = st.file_uploader("Carregar Labeling Specification", type=["docx"])
            if not DOCX_AVAILABLE and rm_file:
                st.error("❌ Biblioteca python-docx não está instalada. Execute 'pip install python-docx' e reinicie a aplicação.")
                rm_file = None
            # Definir rm_file_type para Word quando for Labeling Specification
            rm_file_type = "Word (.docx)"
        
        with col2:
            st.info("📊 Os IDs serão extraídos da aba 'Labeling and Learning Materials', coluna 'PRS ID'")
            prs_file = st.file_uploader("Carregar PRS DOC", type=["xlsx"])
            
        # Definir tipo de documento para uso posterior
        doc_type = "Labeling Specification"


        if rm_file:
            st.warning('⚠️ Arquivo recebido, bloco de upload está sendo executado!')
            
            # Verificar qual tipo de documento está sendo processado
            document_type = "Risk Matrix" if doc_type == "Risk Management Matrix" else "Labeling Specification"
            st.info(f"📄 Processando documento: {document_type}")
            
            # Processar arquivo baseado no tipo selecionado e no documento
            if doc_type == "Labeling Specification":
                # Processar Labeling Specification (sempre em formato Word)
                with st.spinner("🔍 Analisando Labeling Specification..."):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        tmp.write(rm_file.read())
                        rm_path = tmp.name
                    
                    # Usar função especializada para Labeling Specification
                    st.info("🔎 Buscando seção '13. Annex 1 – Matrix of PRS and TCs' no documento...")
                    rm_ids = extract_tasy_ids_from_labeling_spec(rm_path)
            
            # Caso contrário, tratar como Risk Matrix (Excel ou Word)
            elif rm_file_type == "Excel (.xlsx)":
                # Processar arquivo Excel (código original)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                    tmp.write(rm_file.read())
                    rm_path = tmp.name
                
                # Debug: mostrar nomes de abas exatamente como pandas lê
                try:
                    xl = pd.ExcelFile(rm_path)
                    st.info(f"Abas pandas (repr): {[repr(s) for s in xl.sheet_names]}")
                except Exception:
                    pass

                # Debug: mostrar nomes de colunas exatamente como pandas lê, sem header
                try:
                    df_noheader = pd.read_excel(rm_path, sheet_name=0, header=None)
                    with st.expander('🔍 Debug - DataFrame sem header (primeiras 10 linhas)'):
                        st.write(df_noheader.head(10))
                except Exception as e:
                    st.error(f"Erro ao ler DataFrame sem header: {e}")

                # Debug: tentar diferentes skiprows para encontrar o cabeçalho real
                for skip in range(6):
                    try:
                        df_skip = pd.read_excel(rm_path, sheet_name=0, header=0, skiprows=skip)
                        with st.expander(f'🔍 Debug - DataFrame skiprows={skip} (primeiras 10 linhas)'):
                            st.write(df_skip.head(10))
                            st.write(f"Colunas: {[repr(c) for c in df_skip.columns]}")
                    except Exception as e:
                        st.error(f"Erro ao ler DataFrame com skiprows={skip}: {e}")

                # Agora tentar carregar o DataFrame principal (df) normalmente
                try:
                    df = pd.read_excel(rm_path, sheet_name="Risk Management Matrix")
                except Exception:
                    try:
                        df = pd.read_excel(rm_path)
                    except Exception:
                        df = pd.DataFrame()

                # Só executar o restante se df não estiver vazio
                if not df.empty:
                    # Mostrar shape e colunas imediatamente
                    st.info(f"Shape do DataFrame: {df.shape}")
                    st.info(f"Colunas pandas (repr): {[repr(col) for col in df.columns]}")

                    # Debug: DataFrame inteiro e nomes de colunas (repr)
                    with st.expander('🔍 Debug - DataFrame inteiro (primeiras 30 linhas)'):
                        st.write(df.head(30))
                    with st.expander('🔍 Debug - Nomes de colunas (repr)'):
                        st.write([repr(col) for col in df.columns])
                        st.write({i: repr(col) for i, col in enumerate(df.columns)})
                    # Debug: mostrar as 10 primeiras células de todas as colunas
                    with st.expander('🔍 Debug - 10 primeiras células de todas as colunas'):
                        for col in df.columns:
                            st.write(f"Coluna: {repr(col)}")
                            st.write(list(df[col].head(10)))
                    # Debug extra: mostrar até 200 linhas do DataFrame inteiro
                    with st.expander('🔍 Debug - Até 200 linhas do DataFrame inteiro (aba)'):
                        st.write(df.head(200))
                    # Debug: shape do DataFrame
                    with st.expander('🔍 Debug - Shape (linhas, colunas) do DataFrame'):
                        st.write(df.shape)
                    # Mostrar as primeiras linhas da planilha para debug
                    with st.expander('🔍 Debug - Primeiras linhas da planilha lida'):
                        st.write(df.head(10))

                    # Extrair da coluna correta pelo nome exato, a partir da linha 5
                    col_name = 'Risk Control Measures (requirements)'
                    if col_name not in df.columns:
                        st.error(f"Coluna '{col_name}' não encontrada na aba 'Risk Management Matrix'. Colunas disponíveis: {[repr(col) for col in df.columns]}")
                        all_text = ''
                    else:
                        # Debug: mostrar índice real e valores das 15 primeiras linhas da coluna
                        with st.expander(f'🔍 Debug - Índice e valores das 15 primeiras linhas da coluna {col_name}'):
                            for idx, val in zip(df.index[:15], df[col_name].head(15)):
                                st.write(f"Índice pandas: {idx} | Valor: {val}")
                        # Encontrar a primeira linha não vazia
                        first_valid = df[col_name].first_valid_index()
                        st.write(f"Primeira linha não vazia da coluna '{col_name}': índice pandas {first_valid}")
                        # Pegar a partir da primeira linha não vazia
                        col_data = df[col_name].astype(str).loc[first_valid:]
                        all_text = '\n'.join(col_data)
                        with st.expander(f'🔍 Debug - Primeiras 10 linhas da coluna {col_name} (a partir da primeira não vazia) - texto bruto'):
                            for i, (idx, val) in enumerate(zip(col_data.index, col_data.head(10))):
                                st.write(f"Índice pandas: {idx} | Valor: {val}")
                        with st.expander(f'🔍 Debug - Texto concatenado da coluna {col_name} (all_text)'):
                            st.write(all_text[:2000] + ("..." if len(all_text) > 2000 else ""))

                    # Regex mais restritivo: só IDs com 4 grupos numéricos
                    pattern = r'TASY_PRS_ID_\d+\.\d+\.\d+\.\d+'
                    matches = re.findall(pattern, all_text)
                    semicolon_ids = extract_semicolon_ids(all_text)
                    # IDs brutos antes do filtro
                    debug_ids = sorted(list(set(matches).union(semicolon_ids)))
                    with st.expander('🔍 Debug - IDs extraídos antes do filtro'):
                        for id in debug_ids:
                            st.write(id)
                    # Filtro final: só IDs válidos, não problemáticos, não vazios, sem espaços
                    problematic_ids = {"TASY_PRS_ID_1", "TASY_PRS_ID_2", "TASY_PRS_ID_6"}
                    rm_ids = [id.strip() for id in debug_ids if is_valid_tasy_id(id) and id.strip() and id not in problematic_ids]
                    rm_ids = sorted(set(rm_ids))
                    with st.expander('🔍 Debug - IDs finais após filtro'):
                        for id in rm_ids:
                            st.write(id)
                else:
                    st.error("❌ Não foi possível processar o arquivo Excel da Risk Matrix.")
                    rm_ids = []
                    
            else:
                # Processar arquivo Word .docx
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp.write(rm_file.read())
                    rm_path = tmp.name
                
                # Extrair IDs do arquivo Word usando a função especializada com base no tipo de documento
                if doc_type == "Risk Management Matrix":
                    # Processar como Risk Matrix Word
                    rm_ids = extract_tasy_ids_from_docx(rm_path)
                else:
                    # Processar como Labeling Specification
                    st.info("🔎 Analisando documento Labeling Specification...")
                    rm_ids = extract_tasy_ids_from_labeling_spec(rm_path)
            
            # Parte comum: exibir e processar os IDs encontrados (independente do formato do arquivo)
            if rm_ids:
                if doc_type == "Labeling Specification":
                    st.success(f"✅ Encontrados {len(rm_ids)} TASY_PRS_ID únicos no Labeling Specification")
                else:
                    st.success(f"✅ Encontrados {len(rm_ids)} TASY_PRS_ID únicos no arquivo")
                
                df_ids = pd.DataFrame({"TASY_PRS_ID": rm_ids})
                st.dataframe(df_ids, use_container_width=True)
                
                # Definir nome do arquivo com base no tipo de documento
                filename = "tasy_prs_ids_risk_matrix.csv" if doc_type == "Risk Management Matrix" else "tasy_prs_ids_labeling_spec.csv"
                
                csv = df_ids.to_csv(index=False).encode('utf-8')
                st.download_button(
                    "⬇️ Download CSV",
                    csv,
                    filename,
                    "text/csv",
                    key='download-csv'
                )
                
                # Se o PRS DOC também foi carregado, oferecer comparação
                compare_button_text = "Comparar com PRS DOC"
                if doc_type == "Labeling Specification":
                    compare_button_text = "Comparar com PRS DOC (aba 'Labeling and Learning Materials')"
                
                if prs_file and st.button(compare_button_text):
                    with st.spinner("Extraindo TASY_PRS_ID do PRS DOC..."):
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                            tmp.write(prs_file.read())
                            prs_path = tmp.name
                        # Extrair IDs do PRS DOC de acordo com o tipo de documento
                        if doc_type == "Labeling Specification":
                            # Para Labeling Specification, extrair da aba específica
                            prs_ids = extract_tasy_ids_from_prs_doc(prs_path, sheet_name="Labeling and Learning Materials")
                            prs_doc_context = "PRS DOC (aba 'Labeling and Learning Materials')"
                        else:
                            # Para Risk Matrix, extrair da aba padrão
                            prs_ids = extract_tasy_ids_from_prs_doc(prs_path)
                            prs_doc_context = "PRS DOC"
                        
                        if prs_ids:
                            st.success(f"✅ Encontrados {len(prs_ids)} TASY_PRS_ID no {prs_doc_context}")
                            comparison = compare_tasy_ids(rm_ids, prs_ids)
                            st.subheader("Resultados da Comparação")
                            
                            # Definir rótulos com base no tipo de documento
                            doc_label = "Risk Matrix" if doc_type == "Risk Management Matrix" else "Labeling Specification"
                            
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("IDs em comum", len(comparison['common']))
                            with col2:
                                st.metric(f"Somente no {doc_label}", len(comparison['only_in_rm']))
                            with col3:
                                st.metric(f"Somente no {prs_doc_context}", len(comparison['only_in_prs']))
                            
                            tab1, tab2, tab3 = st.tabs([
                                "🔄 IDs em comum", 
                                f"⚠️ Somente no {doc_label}", 
                                f"❓ Somente no {prs_doc_context}"
                            ])
                            
                            with tab1:
                                if comparison['common']:
                                    st.dataframe(pd.DataFrame({"TASY_PRS_ID": comparison['common']}), use_container_width=True)
                                else:
                                    st.warning("Não há IDs em comum!")
                            
                            with tab2:
                                if comparison['only_in_rm']:
                                    st.dataframe(pd.DataFrame({"TASY_PRS_ID": comparison['only_in_rm']}), use_container_width=True)
                                else:
                                    st.success(f"Todos os IDs do {doc_label} estão no PRS DOC!")
                            
                            with tab3:
                                if comparison['only_in_prs']:
                                    st.dataframe(pd.DataFrame({"TASY_PRS_ID": comparison['only_in_prs']}), use_container_width=True)
                                else:
                                    st.success(f"Todos os IDs do {prs_doc_context} estão no {doc_label}!")
                                    
                            # Opção para exportar resultados
                            st.subheader("Exportar Resultados")
                            
                            # Definir rótulos para os status com base no tipo de documento
                            doc_label = "Risk Matrix" if doc_type == "Risk Management Matrix" else "Labeling Specification"
                            
                            results = {
                                "common_ids": comparison['common'],
                                "only_in_rm": comparison['only_in_rm'],
                                "only_in_prs": comparison['only_in_prs']
                            }
                            
                            # Criar um DataFrame para cada categoria
                            df_common = pd.DataFrame({"TASY_PRS_ID": comparison['common'], "Status": "Presente em ambos"})
                            df_only_rm = pd.DataFrame({"TASY_PRS_ID": comparison['only_in_rm'], "Status": f"Somente no {doc_label}"})
                            df_only_prs = pd.DataFrame({"TASY_PRS_ID": comparison['only_in_prs'], "Status": f"Somente no {prs_doc_context}"})
                            
                            # Juntar todos em um único DataFrame
                            df_all = pd.concat([df_common, df_only_rm, df_only_prs])
                            
                            # Definir nome do arquivo com base no tipo de documento
                            filename = f"comparison_results_{doc_type.lower().replace(' ', '_')}_vs_prs_doc.csv"
                            
                            # Opção para baixar como CSV
                            csv_all = df_all.to_csv(index=False).encode('utf-8')
                            st.download_button(
                                "⬇️ Download CSV com Todos os Resultados",
                                csv_all,
                                filename,
                                "text/csv",
                                key='download-csv-all'
                            )
                        else:
                            st.error("❌ Nenhum TASY_PRS_ID encontrado no PRS DOC.")
            else:
                doc_desc = ""
                if doc_type == "Risk Management Matrix":
                    doc_desc = f"Risk Matrix {'Excel' if rm_file_type == 'Excel (.xlsx)' else 'Word'}"
                else:
                    doc_desc = "Labeling Specification"
                    
                st.error(f"❌ Nenhum TASY_PRS_ID encontrado no arquivo {doc_desc}.")